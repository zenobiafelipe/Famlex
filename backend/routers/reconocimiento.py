from fastapi import APIRouter, Form, Depends
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import locale
import os
import unicodedata

from database import get_db
from models import Documento
from routers.auth import obtener_usuario_actual

router = APIRouter()

def normalizar(texto: str) -> str:
    if not texto:
        return ""
    return unicodedata.normalize("NFD", texto.strip().lower()).encode("ascii", "ignore").decode("utf-8")
@router.post("/generar/reconocimiento_paternidad")
async def generar_reconocimiento_paternidad(
    promovente: str = Form(...),
    menor: str = Form(...),
    edad_menor: str = Form(...),
    fecha_nacimiento: str = Form(...),
    direccion: str = Form(...),
    demandado: str = Form(...),
    tipo_relacion: str = Form(...),
    periodo_relacion: str = Form(...),
    motivo: str = Form(...),
    conoce_trabajo: str = Form(...),
    trabajo: str = Form(None),
    direccion_trabajo: str = Form(None),
    ingreso: str = Form(None),
    direccion_demandado: str = Form(...),
    solicita_pension: str = Form(...),
    porcentaje: str = Form(None),
    fecha_incumplimiento: str = Form(None),
    prueba_adn: str = Form(...),
    desea_testigos: str = Form(...),
    testigos: str = Form(None),
    desea_agregar_otros_abogados: str = Form(...),
    abogados: str = Form(None),
    usuario=Depends(obtener_usuario_actual),
    db: Session = Depends(get_db)
):
    promovente = promovente.strip().title()
    menor = menor.strip().title()
    demandado = demandado.strip().title()

    locale.setlocale(locale.LC_TIME, 'es_MX.UTF-8')
    fecha_actual = datetime.datetime.now().strftime("%d de %B de %Y")
    ciudad = "Ciudad de México"
    desea_mas = normalizar(desea_agregar_otros_abogados)
    abogado_lista = [f"{usuario.nombre_completo}:{usuario.cedula_profesional}"]

    if desea_mas == "si" and abogados:
        abogado_lista += [a.strip() for a in abogados.split(";") if a.strip()]

    plural = len(abogado_lista) > 1
    if plural:
        texto_abogados = ", ".join([f"{a.split(':')[0]} (Cédula {a.split(':')[1]})" for a in abogado_lista])
        autorizacion = f"a los C.C. Licenciados en Derecho {texto_abogados}"
    else:
        nombre, cedula = abogado_lista[0].split(":")
        autorizacion = f"al C. Licenciado en Derecho {nombre} (Cédula {cedula})"
#Fecha nacimiento
    try:
        fecha_dt = datetime.datetime.strptime(fecha_nacimiento, "%Y-%m-%d")
        fecha_formateada = fecha_dt.strftime("%d de %B de %Y").capitalize()
    except ValueError:
        fecha_formateada = fecha_nacimiento
#Fecha incumplimiento
    try:
        fecha_incumplimiento_dt = datetime.datetime.strptime(fecha_incumplimiento, "%Y-%m-%d")
        fecha_formateada_incumplimiento = fecha_incumplimiento_dt.strftime("%d de %B de %Y").capitalize()
    except ValueError:
        fecha_formateada_incumplimiento = fecha_incumplimiento

#Creación del documento
    doc = Document()

    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    encabezado.add_run(
        f"{promovente.upper()}\nEn representación de {menor.upper()}\nVs\n{demandado.upper()}\n"
        f"JUICIO: RECONOCIMIENTO DE PATERNIDAD\nEXPEDIENTE: __________\nSECRETARÍA: __________"
    ).font.size = Pt(14)

    doc.add_paragraph("\nC. JUEZ DE LO FAMILIAR EN TURNO\nPODER JUDICIAL DE LA CIUDAD DE MÉXICO\nP R E S E N T E.\n")

    doc.add_paragraph(
        f"{promovente}, por mi propio derecho y en representación de el menor {menor}, quien actualmente tiene {edad_menor} años de edad, "
        f"señalando como domicilio para oír notificaciones {direccion}, autorizando para tales efectos {autorizacion}, comparezco y expongo:\n"
    )

    doc.add_heading("P R E S T A C I O N E S", level=1)
    doc.add_paragraph(f"1. El reconocimiento judicial de paternidad del C. {demandado} a favor de el menor {menor}.")
    if normalizar(solicita_pension) == "si":
        doc.add_paragraph("2. El aseguramiento y fijación de una pensión alimenticia provisional y definitiva.")
    if normalizar(prueba_adn) == "si":
        doc.add_paragraph("3. La práctica de prueba pericial en genética molecular (ADN).")
    doc.add_paragraph("4. El pago de costas procesales.")
    doc.add_heading("H E C H O S", level=1)
    doc.add_paragraph(f"1. La menor {menor} nació el {fecha_formateada}.")
    doc.add_paragraph(f"2. Sostuve una relación de tipo {tipo_relacion} con el demandado durante el periodo {periodo_relacion}.")
    doc.add_paragraph(f"3. Considero que es el padre de la menor porque {motivo}.")

    if normalizar(conoce_trabajo) == "si" and trabajo and direccion_trabajo and ingreso:
        doc.add_paragraph(f"4. El demandado trabaja en {trabajo}, ubicado en {direccion_trabajo}, con un ingreso mensual aproximado de ${ingreso}.")
    else:
        doc.add_paragraph("4. Se desconoce el lugar de trabajo actual del demandado.")

    doc.add_paragraph(f"5. El demandado tiene su domicilio en {direccion_demandado} y no ha reconocido voluntariamente a el menor.")

    if normalizar(solicita_pension) == "si" and porcentaje and fecha_formateada_incumplimiento:
        doc.add_paragraph(f"6. Solicito una pensión del {porcentaje}% sobre sus ingresos, ya que ha incumplido desde {fecha_formateada_incumplimiento}.")
    doc.add_heading("D E R E C H O", level=1)
    doc.add_paragraph(
        "Artículos 4º Constitucional, 361 al 380 y 391 del Código Civil para la Ciudad de México; "
        "255 y siguientes del Código de Procedimientos Civiles local; así como la Convención sobre los Derechos del Niño."
    )

    doc.add_heading("P R U E B A S", level=1)
    num_prueba = 1

    doc.add_paragraph(f"{num_prueba}. Acta de nacimiento de el menor.")
    num_prueba += 1

    if normalizar(desea_testigos) == "si" and testigos:
        doc.add_paragraph(f"{num_prueba}. Testigos: {testigos}.")
        num_prueba += 1

    if normalizar(prueba_adn) == "si":
        doc.add_paragraph(f"{num_prueba}. Prueba pericial en genética molecular (ADN).")
        num_prueba += 1

    doc.add_paragraph(f"{num_prueba}. Presuncional legal y humana.")
    num_prueba += 1

    doc.add_paragraph(f"{num_prueba}. Instrumental de actuaciones.")

    doc.add_heading("P E T I C I O N E S", level=1)
    doc.add_paragraph("PRIMERO. Tenerme por presentado con esta demanda y anexos.")
    doc.add_paragraph("SEGUNDO. Admitir el juicio y ordenar el emplazamiento del demandado.")
    doc.add_paragraph(f"TERCERO. Dictar sentencia que declare la paternidad del C. {demandado} respecto de el menor.")
    if normalizar(solicita_pension) == "si":
        doc.add_paragraph("CUARTO. Fijar y asegurar la pensión alimenticia solicitada.")
    if normalizar(prueba_adn) == "si":
        doc.add_paragraph("QUINTO. Girar oficio para la práctica de prueba pericial en ADN.")
    doc.add_paragraph("Último. Condenar al demandado al pago de costas del juicio.")
    doc.add_paragraph(f"\nPROTESTO LO NECESARIO.\n{ciudad}, a {fecha_actual}\n")
    doc.add_paragraph("_______________________________")
    doc.add_paragraph(promovente.upper())

    excluir_justificacion = [
        "JUICIO: RECONOCIMIENTO DE PATERNIDAD",
        "EXPEDIENTE: __________",
        "SECRETARÍA: __________",
        "C. JUEZ DE LO FAMILIAR EN TURNO",
        "P R E S E N T E:",
        "PROTESTO LO NECESARIO."
    ]

    for p in doc.paragraphs:
        if not any(clave in p.text for clave in excluir_justificacion):
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    nombre_archivo = f"Reconocimiento_Paternidad_{promovente.replace(' ', '_')}.docx"
    carpeta_usuario = f"documentos_usuario/{usuario.id}"
    os.makedirs(carpeta_usuario, exist_ok=True)
    ruta_completa = os.path.join(carpeta_usuario, nombre_archivo)
    doc.save(ruta_completa)

    nuevo_doc = Documento(
        usuario_id=usuario.id,
        nombre=nombre_archivo,
        ruta=ruta_completa
    )
    db.add(nuevo_doc)
    db.commit()

    return FileResponse(
        ruta_completa,
        filename=nombre_archivo,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
